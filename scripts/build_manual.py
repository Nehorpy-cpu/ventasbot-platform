"""Genera el manual ilustrado de VentasBot en formato DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "manual" / "Manual_de_uso_VentasBot.docx"
IMAGES = ROOT / "docs" / "manual" / "images"

NAVY = RGBColor(16, 35, 29)
GREEN = RGBColor(28, 138, 91)
LIME = RGBColor(207, 255, 65)
INK = RGBColor(17, 32, 28)
MUTED = RGBColor(101, 115, 109)
LIGHT = "EAF4EE"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def set_font(run, size=None, bold=None, color=None, name="Aptos") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Heading 1", 20, NAVY, 14, 7),
        ("Heading 2", 14, GREEN, 10, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("VENTASBOT  /  GUÍA OPERATIVA")
    set_font(run, 8, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("VentasBot Platform  ·  Página ")
    set_font(run, 8, False, MUTED)
    add_page_field(footer)


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    set_font(run, 9, True, GREEN)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.7)
    cell = table.cell(0, 0)
    cell.width = Inches(6.7)
    set_cell_fill(cell, LIGHT)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title + "  ")
    set_font(title_run, 10.5, True, GREEN)
    body_run = p.add_run(body)
    set_font(body_run, 10.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc: Document, items: list[str]) -> None:
    numbering = doc.part.numbering_part.element
    list_style = doc.styles["List Number"]
    base_num_id = int(list_style._element.pPr.numPr.numId.val)
    base_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == base_num_id)
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_num_id = max((int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))), default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = next_num_id


def add_figure(doc: Document, filename: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    width = 5.85 if filename == "02-pedidos-desktop.png" else 6.35
    inline = run.add_picture(str(IMAGES / filename), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_font(run, 8.5, False, MUTED)
    run.italic = True


def add_status_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(5.05)
    hdr = table.rows[0].cells
    hdr[0].text = "ESTADO"
    hdr[1].text = "SIGNIFICADO OPERATIVO"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_fill(cell, "10231D")
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            set_font(run, 8.5, True, RGBColor(255, 255, 255))
    for status, meaning in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.65)
        cells[1].width = Inches(5.05)
        cells[0].text = status
        cells[1].text = meaning
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        for run in cells[0].paragraphs[0].runs:
            set_font(run, 9, True, GREEN)
        for run in cells[1].paragraphs[0].runs:
            set_font(run, 9.5, False, INK)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> None:
    doc = Document()
    style_document(doc)

    add_kicker(doc, "Comercio conversacional · Manual de capacitación")
    title = doc.add_paragraph(style="Title")
    title.add_run("Manual de uso\nVentasBot Platform")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("De la conversación de WhatsApp a la entrega, en una sola operación.")
    set_font(run, 14, False, MUTED)
    add_figure(doc, "01-resumen-desktop.png", "Vista general de la empresa Pizzería Demo con historial precargado.")
    add_callout(doc, "ALCANCE DE ESTA GUÍA", "Recorrido funcional para dueños, vendedores, depósito, despacho y delivery. Todos los datos mostrados son ficticios.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versión demo · Agosto de 2026")
    set_font(run, 9, True, MUTED)

    page_break(doc)
    doc.add_heading("1. Empezar en tres minutos", level=1)
    doc.add_heading("Encender la aplicación", level=2)
    doc.add_paragraph("Desde PowerShell, dentro de la carpeta del proyecto, ejecutá:")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_after = Pt(10)
    run = code.add_run(".\\.venv-clean\\Scripts\\python.exe -m app.seed\n.\\.venv-clean\\Scripts\\python.exe -m uvicorn app.main:app --reload")
    set_font(run, 9.5, False, NAVY, "Consolas")
    doc.add_heading("Ingresar como empresa demo", level=2)
    add_steps(doc, [
        "Abrí http://127.0.0.1:8000/panel/.",
        "Empresa: pizzeria-demo.",
        "Correo: demo@ventasbot.local.",
        "Contraseña: VentasBotDemo2026!.",
        "Presioná Entrar al panel.",
    ])
    add_callout(doc, "SEGURIDAD", "Las credenciales de esta guía son solo para la computadora local. Cambialas antes de publicar o conectar servicios reales.")
    doc.add_heading("Qué contiene la demo", level=2)
    add_bullets(doc, [
        "3 productos y stock de ejemplo.",
        "6 clientes con conversaciones de WhatsApp simuladas.",
        "5 pedidos: pago pendiente, preparando, en tránsito, entregado y cancelado.",
        "5 pagos, 5 facturas y 2 entregas con eventos de tracking.",
    ])

    page_break(doc)
    doc.add_heading("2. Leer el resumen", level=1)
    doc.add_paragraph("El Resumen es la vista rápida del dueño o responsable de turno. Muestra el tamaño del catálogo, la cantidad de pedidos y el valor total operado.")
    add_figure(doc, "01-resumen-desktop.png", "Resumen operativo: productos, pedidos y valor acumulado no cancelado.")
    add_bullets(doc, [
        "Productos: artículos disponibles en el tenant actual.",
        "Pedidos: ventas registradas para esta empresa.",
        "Valor operado: suma de pedidos no cancelados; no necesariamente equivale a cobros confirmados.",
    ])

    page_break(doc)
    doc.add_heading("3. Atender conversaciones", level=1)
    doc.add_paragraph("La bandeja reúne el historial de WhatsApp. Elegí un cliente a la izquierda para revisar mensajes y responder desde el compositor inferior.")
    add_figure(doc, "03-conversaciones-desktop.png", "Bandeja de conversaciones con mensajes entrantes y respuestas del bot.")
    add_status_table(doc, [
        ("BOT", "La automatización conduce la compra y recopila los datos."),
        ("HUMAN", "Un operador tomó el control de la atención."),
        ("CLOSED", "La conversación terminó y queda disponible como historial."),
    ])
    add_callout(doc, "MODO DEMO", "Sin Meta conectado, una respuesta se registra en el CRM pero no llega a un número real de WhatsApp.")

    page_break(doc)
    doc.add_heading("4. Gestionar pedidos", level=1)
    doc.add_paragraph("Pedidos concentra cada venta y su avance. La columna Acción habilita solamente la siguiente operación válida para evitar saltos de estado.")
    add_figure(doc, "02-pedidos-desktop.png", "Cinco escenarios de pedido preparados para practicar la operación.")
    add_status_table(doc, [
        ("PENDING_PAYMENT", "Espera tarjeta, transferencia o comprobante."),
        ("CONFIRMED", "Venta aceptada y asegurada."),
        ("PREPARING", "Depósito o cocina está preparando."),
        ("READY", "Pedido terminado y listo para despacho."),
        ("IN_TRANSIT", "Delivery va hacia el cliente."),
        ("DELIVERED", "Entrega completada."),
        ("CANCELLED", "Pedido cerrado; el stock reservado se libera cuando corresponde."),
    ])

    page_break(doc)
    doc.add_heading("5. Revisar facturación", level=1)
    doc.add_paragraph("Facturación almacena nombre o razón social, RUC, importe, estado y referencia del documento vinculado al pedido.")
    add_figure(doc, "04-facturacion-desktop.png", "Registros fiscales demo con estados pendiente, emitido y cancelado.")
    add_status_table(doc, [
        ("PENDING", "Datos recopilados; emisión todavía pendiente."),
        ("ISSUED", "Documento generado con número y, en producción, enlace KuDE."),
        ("CANCELLED", "Documento anulado o asociado a una venta cancelada."),
    ])
    add_callout(doc, "SIFEN", "La demo no emite comprobantes tributarios reales. Producción requiere homologación, certificado y credenciales del contribuyente.")

    page_break(doc)
    doc.add_heading("6. Coordinar entregas", level=1)
    doc.add_paragraph("Entregas permite al despacho y al delivery revisar cliente, horario y estado. Cada cambio agrega un evento al tracking del pedido.")
    add_figure(doc, "05-entregas-desktop.png", "Una entrega completada y otra en tránsito listas para demostración.")
    add_status_table(doc, [
        ("ASSIGNED", "Repartidor designado."),
        ("PICKED_UP", "Pedido retirado del depósito."),
        ("IN_TRANSIT", "En camino al cliente."),
        ("ARRIVED", "Delivery llegó al destino."),
        ("DELIVERED", "Entrega finalizada."),
        ("FAILED", "No fue posible completar la entrega."),
    ])

    page_break(doc)
    doc.add_heading("7. Configurar cada empresa", level=1)
    doc.add_paragraph("Cada tenant conserva sus propias integraciones. Las pantallas guardan referencias a variables seguras, no claves privadas en texto visible.")
    add_figure(doc, "06-configuracion-desktop.png", "Configuración independiente de Meta, Bancard y métodos de pago.")
    add_bullets(doc, [
        "Meta: Phone Number ID, WABA, versión Graph y variable del token.",
        "Bancard: comercio, sucursal, retornos y variables de claves.",
        "Transferencia y contra entrega: instrucciones que el bot ofrece al cliente.",
        "Catálogo: sincronización desde Meta, web, WhatsApp Business, CSV o API.",
    ])
    add_callout(doc, "NO PEGAR SECRETOS", "Tokens, claves privadas y certificados se cargan como variables de entorno. Nunca deben copiarse en chats, capturas o documentación.")

    page_break(doc)
    doc.add_heading("8. Roles y responsabilidades", level=1)
    add_status_table(doc, [
        ("PLATFORM_ADMIN", "Administra empresas, demos y operación global."),
        ("TENANT_OWNER", "Controla usuarios, catálogo e integraciones de su empresa."),
        ("TENANT_MANAGER", "Supervisa la operación diaria."),
        ("SELLER", "Atiende conversaciones y ventas."),
        ("WAREHOUSE", "Gestiona preparación y stock."),
        ("DISPATCHER", "Asigna y supervisa entregas."),
        ("DRIVER", "Actualiza el recorrido de pedidos asignados."),
    ])
    doc.add_heading("Recorrido práctico", level=2)
    add_steps(doc, [
        "Confirmá que el Resumen muestre cinco pedidos.",
        "Compará una conversación BOT con una CLOSED.",
        "En Pedidos, avanzá el registro PREPARING a READY.",
        "En Facturación, distinguí PENDING, ISSUED y CANCELLED.",
        "En Entregas, avanzá IN_TRANSIT a ARRIVED.",
        "Revisá Configuración sin agregar credenciales reales todavía.",
    ])

    page_break(doc)
    doc.add_heading("9. Crear capacitación con NotebookLM", level=1)
    doc.add_paragraph("El archivo MANUAL_NOTEBOOKLM.md contiene el mismo conocimiento en texto estructurado y es ideal como fuente para consultas, resúmenes y materiales de capacitación.")
    add_steps(doc, [
        "Creá un cuaderno nuevo en Google NotebookLM.",
        "Elegí Añadir fuente y subí MANUAL_NOTEBOOKLM.md o este manual Word.",
        "Agregá README.md, ROADMAP.md y SECURITY.md si querés incluir arquitectura y seguridad.",
        "Pedí respuestas basadas únicamente en las fuentes cargadas.",
    ])
    doc.add_heading("Prompts sugeridos", level=2)
    add_bullets(doc, [
        "Creá una guía de capacitación separada por roles.",
        "Explicá el recorrido desde WhatsApp hasta la entrega.",
        "Prepará un checklist diario para depósito y despacho.",
        "Indicá qué funciones son demo y cuáles requieren homologación.",
        "Generá preguntas frecuentes para incorporar una empresa nueva.",
    ])
    add_callout(doc, "FUENTES RECOMENDADAS", "Manual, README, Roadmap y guía de Seguridad. No subas el archivo .env porque contiene credenciales locales.")

    page_break(doc)
    doc.add_heading("10. Límites antes de producción", level=1)
    add_bullets(doc, [
        "Meta requiere número verificado, app configurada, webhook HTTPS y token seguro.",
        "Bancard requiere comercio habilitado, ambiente homologado y callbacks autorizados.",
        "SIFEN requiere homologación y certificados tributarios del contribuyente.",
        "SQLite es para prueba local; producción debe usar PostgreSQL, backups y monitoreo.",
        "El despliegue público debe usar HTTPS, secretos externos y controles de acceso.",
    ])
    add_callout(doc, "PRÓXIMO PASO", "Después de validar el recorrido demo, se puede conectar una empresa piloto y probar Meta/Bancard en sandbox antes de habilitar transacciones reales.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VentasBot Platform\nTu negocio vende. VentasBot coordina.")
    set_font(run, 16, True, NAVY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
